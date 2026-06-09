#!/usr/bin/env python3
"""
Convert Markdown file(s) sang .docx (Word).

Requires: pip install python-docx markdown beautifulsoup4

Usage:
  # Single file
  python tools/md_to_docx.py --file testing-output/test-plan/Test_Plan_FDP_Sprint12_v1.0.0.md

  # Toàn bộ folder
  python tools/md_to_docx.py --folder testing-output/test-plan

  # Chỉ định output dir
  python tools/md_to_docx.py --folder testing-output --out testing-output/docx --recursive
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import markdown
    from bs4 import BeautifulSoup, Tag
except ImportError:
    print("ERROR: Missing dependencies. Run:")
    print("  pip install python-docx markdown beautifulsoup4")
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent

TABLE_HEADER_HEX = "D5E8F0"


# ── Docx helpers ──────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str = TABLE_HEADER_HEX) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _apply_inline(run, text: str, bold=False, italic=False, code=False) -> None:
    run.text = text
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def _add_paragraph_from_tag(doc: Document, tag: Tag, style: str = "Normal") -> None:
    p = doc.add_paragraph(style=style)
    _fill_paragraph(p, tag)


def _fill_paragraph(p, tag: Tag) -> None:
    for child in tag.children:
        if isinstance(child, str):
            if child.strip():
                p.add_run(child)
        elif child.name in ("strong", "b"):
            p.add_run(child.get_text()).bold = True
        elif child.name in ("em", "i"):
            p.add_run(child.get_text()).italic = True
        elif child.name == "code":
            run = p.add_run(child.get_text())
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif child.name == "a":
            p.add_run(child.get_text())
        elif child.name == "br":
            p.add_run("\n")
        else:
            p.add_run(child.get_text())


def _add_table(doc: Document, tag: Tag) -> None:
    rows = tag.find_all("tr")
    if not rows:
        return
    cols = max(len(r.find_all(["th", "td"])) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for c_idx, cell in enumerate(cells):
            if c_idx >= cols:
                break
            tc = table.rows[r_idx].cells[c_idx]
            tc.text = ""
            p = tc.paragraphs[0]
            _fill_paragraph(p, cell)
            is_header = cell.name == "th" or r_idx == 0
            if is_header:
                p.runs[0].bold = True if p.runs else None
                _set_cell_bg(tc)


def md_to_docx(md_path: Path, out_path: Path) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    html = markdown.markdown(
        md_text,
        extensions=["tables", "fenced_code", "sane_lists"],
    )
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)

    heading_map = {"h1": "Heading 1", "h2": "Heading 2", "h3": "Heading 3",
                   "h4": "Heading 4", "h5": "Heading 5", "h6": "Heading 6"}

    for tag in soup.children:
        if not isinstance(tag, Tag):
            continue
        name = tag.name

        if name in heading_map:
            p = doc.add_paragraph(tag.get_text(), style=heading_map[name])

        elif name == "p":
            _add_paragraph_from_tag(doc, tag)

        elif name in ("ul", "ol"):
            for li in tag.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Bullet" if name == "ul" else "List Number")
                _fill_paragraph(p, li)

        elif name == "table":
            _add_table(doc, tag)

        elif name == "pre":
            code_tag = tag.find("code")
            text = code_tag.get_text() if code_tag else tag.get_text()
            p = doc.add_paragraph(style="Normal")
            run = p.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)

        elif name == "blockquote":
            p = doc.add_paragraph(style="Quote")
            _fill_paragraph(p, tag)

        elif name == "hr":
            doc.add_paragraph("─" * 60)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


# ── CLI ────────────────────────────────────────────────────────────────────────

def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__,
                                     formatter_class=argparse.RawDescriptionHelpFormatter)
    src = parser.add_mutually_exclusive_group(required=True)
    src.add_argument("--file",   help="Single .md file")
    src.add_argument("--folder", help="Folder chứa .md files")
    parser.add_argument("--out",       help="Output folder (default: cùng thư mục với file gốc)")
    parser.add_argument("--recursive", action="store_true", help="Tìm .md trong subfolder")
    args = parser.parse_args()

    if args.file:
        md_files = [Path(args.file)]
    else:
        folder = Path(args.folder)
        if not folder.is_absolute():
            folder = ROOT / folder
        md_files = sorted(folder.rglob("*.md") if args.recursive else folder.glob("*.md"))

    if not md_files:
        print("No .md files found.")
        return 0

    print(f"Converting {len(md_files)} file(s)...\n")
    errors = 0
    for md_path in md_files:
        if not md_path.is_absolute():
            md_path = ROOT / md_path
        if not md_path.exists():
            print(f"  NOT FOUND: {md_path}")
            errors += 1
            continue

        out_dir = Path(args.out) if args.out else md_path.parent
        if not out_dir.is_absolute():
            out_dir = ROOT / out_dir
        out_path = out_dir / md_path.with_suffix(".docx").name

        try:
            md_to_docx(md_path, out_path)
            rel = out_path.relative_to(ROOT) if out_path.is_relative_to(ROOT) else out_path
            print(f"  OK: {rel}")
        except Exception as e:
            print(f"  ERROR: {md_path.name} — {e}")
            errors += 1

    return 0 if errors == 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
